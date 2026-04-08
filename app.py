"""
TalentIQ — AI-Based Resume Screening System
Enhanced Flask backend with:
  - PDF/DOCX resume parsing (pdfplumber, PyPDF2, python-docx)
  - TF-IDF cosine similarity scoring
  - HR login / session auth (SQLite persistent database)
  - Weighted multi-dimensional ranking
  - Model evaluation metrics API
  - Admin user management (add/edit/delete users)
"""

from flask import Flask, render_template, request, jsonify, session, redirect, url_for
import os, math, re, io, hashlib, sqlite3, json
from functools import wraps
from collections import Counter
from datetime import datetime, timedelta, timezone

# Optional heavy parsers
try:
    import pdfplumber
    PDF_ENGINE = "pdfplumber"
except ImportError:
    pdfplumber = None
    try:
        from PyPDF2 import PdfReader
        PDF_ENGINE = "PyPDF2"
    except ImportError:
        PdfReader = None
        PDF_ENGINE = None

try:
    from docx import Document as DocxDocument
    DOCX_ENGINE = "python-docx"
except ImportError:
    DocxDocument = None
    DOCX_ENGINE = None

app = Flask(__name__)
app.secret_key = os.environ.get("TALENTIQ_SECRET", "talentiq-hr-secret-2024")

#  DATABASE SETUP (SQLite) 

DB_PATH = os.path.join(os.path.dirname(__file__), "talentiq_users.db")

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    """Create users table and seed default admin if empty."""
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            name      TEXT    NOT NULL,
            email     TEXT    NOT NULL UNIQUE,
            password  TEXT    NOT NULL,
            role      TEXT    NOT NULL DEFAULT 'hr',
            created_at TEXT   NOT NULL DEFAULT (datetime('now')),
            last_login TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS login_logs (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            email      TEXT    NOT NULL,
            user_name  TEXT,
            role       TEXT,
            status     TEXT    NOT NULL,
            ip_address TEXT,
            logged_at  TEXT    NOT NULL DEFAULT (datetime('now','localtime'))
        )
    """)
    # Seed default users if table is empty
    count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    if count == 0:
        defaults = [
            ("Admin HR",   "admin@talentiq.com",     hash_pw("Admin@123"),     "admin"),
            ("HR Manager", "hr@talentiq.com",         hash_pw("Hr@123"),        "hr"),
            ("Recruiter",  "recruiter@talentiq.com",  hash_pw("Recruiter@123"), "recruiter"),
        ]
        conn.executemany("INSERT INTO users (name,email,password,role) VALUES (?,?,?,?)", defaults)
    conn.commit()
    conn.close()

def hash_pw(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def resolve_user(email, pw, ip=None):
    """Look up user in DB, validate password, and log the attempt."""
    email = email.strip().lower()
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE LOWER(email)=?", (email,)).fetchone()
    # Calculate IST correctly (UTC + 5 hours 30 mins)
    now = (datetime.now(timezone.utc) + timedelta(hours=5, minutes=30)).strftime("%Y-%m-%d %H:%M:%S")
    if row and row["password"] == hash_pw(pw):
        conn.execute("UPDATE users SET last_login=? WHERE id=?", (now, row["id"]))
        conn.execute(
            "INSERT INTO login_logs (email, user_name, role, status, ip_address, logged_at) VALUES (?,?,?,?,?,?)",
            (email, row["name"], row["role"], "success", ip, now)
        )
        conn.commit()
        conn.close()
        return {"id": row["id"], "email": row["email"], "name": row["name"], "role": row["role"]}
    else:
        # Log failed attempt
        name = row["name"] if row else None
        role = row["role"] if row else None
        conn.execute(
            "INSERT INTO login_logs (email, user_name, role, status, ip_address, logged_at) VALUES (?,?,?,?,?,?)",
            (email, name, role, "failed", ip, now)
        )
        conn.commit()
        conn.close()
        return None

# Admin-only decorator
def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        u = session.get("user")
        if not u or u.get("role") != "admin":
            return jsonify({"error": "Admin access required"}), 403
        return f(*args, **kwargs)
    return decorated

# Initialise DB on startup
init_db()

STOP_WORDS = set([
    "with","that","this","have","from","they","will","your","more","been","were",
    "their","than","also","into","about","when","which","there","these","those",
    "would","could","should","must","just","some","each","such","both","only",
    "very","well","what","over","after","before","during","through","between",
    "against","without","within","using","used","work","working","worked","strong",
    "good","great","excellent","experience","years","year","and","the","for","are",
    "was","has","had","not","but","can","you","all","any","its","our","may","who",
    "she","him","her","they","then","than","a","an","in","on","at","to","of","is",
    "be","by","as","or","we","it","do","if","so","up","no","my","he","me","us",
])

SKILL_SYNONYMS = {
    "javascript": ["js","node","nodejs","react","vue","angular","typescript","ts"],
    "python": ["py","django","flask","fastapi","pandas","numpy","scipy"],
    "machine learning": ["ml","ai","deep learning","neural network","tensorflow","pytorch","keras","sklearn","scikit"],
    "sql": ["mysql","postgresql","postgres","sqlite","database","db","oracle","mssql"],
    "cloud": ["aws","azure","gcp","google cloud","amazon web services"],
    "devops": ["ci/cd","docker","kubernetes","k8s","terraform","jenkins","github actions"],
    "react": ["reactjs","react.js","jsx","redux","zustand","next.js","nextjs"],
    "java": ["spring","springboot","maven","gradle","jvm"],
    "restful": ["rest","api","rest api","graphql","grpc"],
    "agile": ["scrum","kanban","jira","sprint"],
}

TECH_PATTERN = re.compile(
    r'\b(react|angular|vue|python|java|typescript|javascript|sql|nosql|aws|azure|gcp|'
    r'docker|kubernetes|node|spring|django|flask|graphql|rest|api|ml|ai|tensorflow|'
    r'pytorch|sklearn|redis|kafka|git|postgres|mysql|mongodb|elasticsearch|linux|'
    r'css|html|scss|redux|zustand|jest|cypress|nextjs|tailwind|swift|kotlin|'
    r'flutter|rust|go|cpp|csharp|dotnet|php|ruby|rails|spark|tableau|powerbi|figma|'
    r'typescript|webpack|vite|rollup|storybook|pyspark|hadoop|airflow|'
    r'fastapi|celery|rabbitmq|nginx|ansible|prometheus|grafana)\b',
    re.IGNORECASE
)

def tokenize(text):
    text = re.sub(r'[^a-z0-9\s]', ' ', text.lower())
    return [t for t in text.split() if len(t) > 1 and t not in STOP_WORDS]
def ngrams(tokens, n): return [" ".join(tokens[i:i+n]) for i in range(len(tokens)-n+1)]
def expand_term(term):
    t = term.lower(); exps = {t}
    for canonical, syns in SKILL_SYNONYMS.items():
        if t == canonical or t in syns:
            exps.add(canonical); exps.update(syns)
    return exps
def semantic_match(jd_term, resume_text):
    rl = resume_text.lower()
    return any(v in rl for v in expand_term(jd_term))

# PDF parsing
def parse_pdf_text(fb):
    if pdfplumber:
        try:
            with pdfplumber.open(io.BytesIO(fb)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages).strip()
        except: pass
    if 'PdfReader' in dir():
        try:
            r = PdfReader(io.BytesIO(fb))
            return "".join(p.extract_text() or "" for p in r.pages).strip()
        except: pass
    return ""

def parse_docx_text(fb):
    if DocxDocument:
        try:
            doc = DocxDocument(io.BytesIO(fb))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text.strip())
            return "\n".join(parts)
        except: pass
    return ""

def parse_resume_file(fb, filename):
    ext = filename.rsplit(".",1)[-1].lower() if "." in filename else "txt"
    if ext == "pdf": raw = parse_pdf_text(fb); eng = PDF_ENGINE or "none"
    elif ext in ("docx","doc"): raw = parse_docx_text(fb); eng = DOCX_ENGINE or "none"
    else:
        try: raw = fb.decode("utf-8","replace")
        except: raw = ""
        eng = "text"
    s = extract_resume_structure(raw)
    s.update({"raw": raw, "engine": eng, "filename": filename, "fileType": ext})
    return s

def extract_resume_structure(text):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    name = ""
    for line in lines[:8]:
        clean = line.split("|")[0].split("—")[0].split(",")[0].strip()
        if "@" in clean or len(clean) > 60 or len(clean) < 3: continue
        # Match: normal case (Riddhi Vegad), all caps (RIDDHI VEGAD), with middle initial (RIDDHI R. VEGAD)
        if re.match(r'^[A-Z][A-Za-z]+[\s.]+[A-Z][A-Za-z]*[\s.]*[A-Za-z]*$', clean):
            # skip obvious non-names
            if not any(w in clean.lower() for w in ['resume','curriculum','vitae','page','address','phone','email','www','http']):
                name = clean.title(); break
    email_m = re.search(r'[\w.+\-]+@[\w\-]+\.[a-z]{2,}', text, re.I)
    phone_m = re.search(r'[\+\d][\d\s\-\(\)]{8,14}\d', text)
    sec_hdrs = {
        "skills": re.compile(r'^(skills?|technical skills?|core skills?|tech stack|technologies|tools|competencies)',re.I),
        "experience": re.compile(r'^(experience|work experience|employment|professional experience)',re.I),
        "education": re.compile(r'^(education|academic|qualification)',re.I),
        "summary": re.compile(r'^(summary|profile|objective|about me)',re.I),
        "certifications": re.compile(r'^(certifications?|certificates?|achievements?|awards?)',re.I),
        "projects": re.compile(r'^(projects?|portfolio|open source)',re.I),
    }
    cur = "other"; slines = {"other": []}
    for line in lines:
        matched = False
        for sec, pat in sec_hdrs.items():
            if pat.match(line) and len(line)<60:
                cur=sec; slines.setdefault(sec,[]); matched=True; break
        if not matched: slines.setdefault(cur,[]); slines[cur].append(line)
    skills_text = " ".join(slines.get("skills",[]))
    raw_skills = re.split(r'[,;|\n]', skills_text)
    extracted_skills = list({s.strip() for s in raw_skills if 1<len(s.strip())<40 and re.search(r'[a-zA-Z]',s)})
    year_matches = re.findall(r'(\d{4})\s*[–\-—to]+\s*(present|\d{4})', text, re.I)
    total_exp = sum(max(0,(2024 if e.lower()=="present" else int(e))-int(s)) for s,e in year_matches)
    edu_text = " ".join(slines.get("education",[]))
    deg_m = re.search(r'(b\.?tech|b\.?e|b\.?s|m\.?tech|m\.?s|m\.?b\.?a|ph\.?d|bachelor|master|doctorate)',edu_text,re.I)
    return {
        "name": name or "Unknown",
        "email": email_m.group(0) if email_m else "",
        "phone": phone_m.group(0) if phone_m else "",
        "skills": extracted_skills,
        "experienceYears": total_exp,
        "degree": deg_m.group(0) if deg_m else "",
        "sections": [k for k,v in slines.items() if v],
    }

def tfidf_cosine(doc_a, doc_b):
    ta = tokenize(doc_a) + ngrams(tokenize(doc_a), 2)
    tb = tokenize(doc_b) + ngrams(tokenize(doc_b), 2)
    tfa, tfb = Counter(ta), Counter(tb)
    vocab = set(tfa)|set(tfb)
    def idf(t): return math.log((3)/((1 if t in tfa else 0)+(1 if t in tfb else 0)+1))+1
    va = {t: (tfa.get(t,0)/max(len(ta),1))*idf(t) for t in vocab}
    vb = {t: (tfb.get(t,0)/max(len(tb),1))*idf(t) for t in vocab}
    dot = sum(va[t]*vb[t] for t in vocab)
    ma = math.sqrt(sum(v*v for v in va.values()))
    mb = math.sqrt(sum(v*v for v in vb.values()))
    if ma==0 or mb==0: return 0.0
    return dot/(ma*mb)

def analyze_jd(jd_text):
    lines = jd_text.split("\n"); req_lines=[]; nice_lines=[]; in_nice=False
    for line in lines:
        if re.search(r'nice[\s\-]to[\s\-]have|bonus|preferred|plus|advantage',line,re.I): in_nice=True
        elif re.search(r'required|must|requirement|responsibilities',line,re.I): in_nice=False
        (nice_lines if in_nice else req_lines).append(line)
    tech = list({t.lower() for t in TECH_PATTERN.findall(jd_text)})
    exp_m = re.search(r'(\d+)\+?\s*years?\s*(of)?\s*(experience|exp)',jd_text,re.I)
    jd_tokens = tokenize(jd_text)
    kws = list({*ngrams(jd_tokens,2),*tech,*[t for t in jd_tokens if len(t)>3]})[:50]
    return {"keywords":kws,"techSkills":tech,"requiredYears":int(exp_m.group(1)) if exp_m else 0,
            "requiredSection":" ".join(req_lines),"niceSection":" ".join(nice_lines)}

def deep_match(resume_text, resume_name, jd):
    jda = analyze_jd(jd)
    rp = extract_resume_structure(resume_text)
    rl = resume_text.lower()
    matched_tech = [t for t in jda["techSkills"] if semantic_match(t,resume_text)]
    missing_tech = [t for t in jda["techSkills"] if not semantic_match(t,resume_text)]
    matched_kw   = [k for k in jda["keywords"]   if semantic_match(k,resume_text)]
    missing_kw   = [k for k in jda["keywords"]   if not semantic_match(k,resume_text)]
    cosine = tfidf_cosine(jd, resume_text)
    cosine_score = min(100, round(cosine*100*2.2))
    tc = len(jda["techSkills"])
    tech_score = round((len(matched_tech)/tc)*100) if tc>0 else round((len(matched_kw)/max(len(jda["keywords"]),1))*100)
    req_yr = jda["requiredYears"]; found_yr = rp["experienceYears"]
    if req_yr>0 and found_yr>0: exp_score = min(100, round((found_yr/req_yr)*80))
    elif re.search(r'senior|lead|principal|architect|manager',resume_text,re.I): exp_score=75
    elif re.search(r'junior|intern|entry',resume_text,re.I): exp_score=35
    else: exp_score=50
    role_words = [w for w in re.split(r'\W+',jd.lower()) if len(w)>4][:12]
    role_score = min(100,round((sum(1 for w in role_words if w in rl)/max(len(role_words),1))*100)+10)
    soft_jd = bool(re.search(r'lead|mentor|manag|team|communication|collaboration',jd,re.I))
    soft_rs = bool(re.search(r'lead|mentor|manag|team|communication|collaboration',resume_text,re.I))
    lead_score = (75 if soft_rs else 30) if soft_jd else 60
    overall = min(97,max(15,round(tech_score*0.40+cosine_score*0.20+exp_score*0.20+role_score*0.12+lead_score*0.08)))
    strengths=[]; gaps=[]
    if matched_tech: strengths.append(f"Matched {len(matched_tech)}/{tc} required skills: {', '.join(matched_tech[:5])}")
    if found_yr>=req_yr>0: strengths.append(f"Experience meets requirement: {found_yr} yrs (required {req_yr}+)")
    if re.search(r'mentor|lead|senior',resume_text,re.I): strengths.append("Leadership & mentoring experience evident")
    if re.search(r'\d+%|reduced|improved|increased|delivered',resume_text,re.I): strengths.append("Quantified impact metrics present")
    if cosine_score>=60: strengths.append(f"High semantic similarity to JD (TF-IDF: {cosine_score}/100)")
    while len(strengths)<2: strengths.append("Keyword alignment with core requirements")
    if missing_tech: gaps.append(f"Missing skills: {', '.join(missing_tech[:5])}")
    if req_yr>0 and found_yr<req_yr: gaps.append(f"Experience gap: {found_yr} yrs found, {req_yr}+ required")
    if len(missing_tech)>3: gaps.append(f"{len(missing_tech)} required technologies not found")
    label = ("strong match for this role and should be prioritised for interview" if overall>=80
             else "good fit with a few gaps worth exploring" if overall>=60
             else "partial match — some key requirements are missing" if overall>=40
             else "limited alignment with this role's core requirements")
    has_sec = bool(re.search(r'experience|education|skills',resume_text,re.I))
    has_bul = bool(re.search(r'[-•*]',resume_text))
    has_con = bool(rp["email"] or re.search(r'\d{10}|\+\d',resume_text))
    has_sum = bool(re.search(r'summary|objective|profile',resume_text,re.I))
    kw_pct  = round((len(matched_tech)/max(tc,1))*100)
    ats = min(96,max(25,round((20 if has_sec else 0)+(15 if has_con else 0)+(10 if has_bul else 0)+(10 if has_sum else 0)+kw_pct*0.45)))
    return {
        "match": {
            "score": overall,
            "cosineSimilarity": cosine_score,
            "dimensions": {"Technical Skills":tech_score,"TF-IDF Similarity":cosine_score,"Experience Level":exp_score,"Role Alignment":role_score,"Leadership & Soft":lead_score},
            "strengths": strengths[:4],
            "gaps": gaps[:3],
            "recommendation": f"{resume_name} is a {label}. {'Recommend technical screen.' if overall>=60 else 'Consider if pool is limited.'}",
            "explanation": {"matchedSkills":matched_tech[:10],"missingSkills":missing_tech[:8],"matchedKeywords":matched_kw[:12],"missingKeywords":missing_kw[:8],"experienceYears":found_yr,"requiredYears":req_yr,"totalJDSkills":tc,"matchedTechCount":len(matched_tech),"cosineSimilarity":cosine_score},
            "parsedResume": rp,
        },
        "ats": {
            "atsScore": ats,
            "summary": f"Resume scores {ats}/100 on ATS compatibility.",
            "subScores": {"Keyword Optimization":min(100,round(kw_pct*1.05)),"Formatting & Structure":min(100,60+(20 if has_bul else 0)+(10 if has_sum else 0)) if has_sec else 35,"Section Completeness":min(100,55+(15 if has_sum else 0)+(10 if has_con else 0)) if has_sec else 30,"Readability":78 if has_bul else 50,"Contact Info":85 if has_con else 40},
            "sections": {"Contact Information":has_con,"Professional Summary":has_sum,"Work Experience":bool(re.search(r'experience|work|employment',resume_text,re.I)),"Education":bool(re.search(r'education|university|degree|college',resume_text,re.I)),"Skills Section":bool(re.search(r'skills|technologies|tools',resume_text,re.I)),"Certifications / Achievements":bool(re.search(r'certif|award|achievement',resume_text,re.I))},
            "keywordMatchPct": kw_pct,
            "keywords": {"matched":matched_tech[:10],"missing":missing_tech[:8]},
            "formattingIssues": [x for x in [not has_sec and "Use standard section headers",not has_bul and "Use bullet points",not has_con and "Add contact information",not has_sum and "Add Professional Summary"] if x][:4],
            "improvements": [x for x in [missing_tech and f"Add missing keywords: {', '.join(missing_tech[:4])}","Quantify achievements","Use standard section headers","Tailor resume to mirror JD phrases"] if x][:5],
        }
    }

def compute_evaluation_metrics(candidates):
    total = len(candidates)
    if total==0: return {"error":"No candidates"}
    scores = [c.get("score",0) for c in candidates]
    avg = round(sum(scores)/total)
    ats_scores = [c.get("ats",{}).get("atsScore",0) for c in candidates if c.get("ats")]
    avg_ats = round(sum(ats_scores)/len(ats_scores)) if ats_scores else 0
    hi = len([c for c in candidates if c.get("score",0)>=80])
    def is_qual(c):
        expl = (c.get("analysis") or c.get("match",{})).get("explanation",{})
        m=expl.get("matchedTechCount",0); t=expl.get("totalJDSkills",1) or 1
        return (m/t)>=0.5 or (c.get("ats") or {}).get("atsScore",0)>=65
    qual = [c for c in candidates if is_qual(c)]; qs={id(c) for c in qual}
    sh   = [c for c in candidates if c.get("score",0)>=80]
    TP=sum(1 for c in sh if id(c) in qs); FP=sum(1 for c in sh if id(c) not in qs)
    FN=sum(1 for c in candidates if c.get("score",0)<80 and id(c) in qs)
    TN=sum(1 for c in candidates if c.get("score",0)<80 and id(c) not in qs)
    prec = min(99,round(TP/hi*100)) if hi>0 else 0
    ws=[c for c in candidates if (c.get("analysis") or {}).get("explanation",{}).get("matchedTechCount",0)>0]
    ct=[c for c in ws if c.get("score",0)>=60]
    rec = min(99,round(len(ct)/len(ws)*100)) if ws else 75
    f1  = round(2*prec*rec/(prec+rec)) if prec+rec>0 else 0
    acc = min(99,round((TP+TN)/total*100))
    var = sum((s-avg)**2 for s in scores)/max(total,1)
    stddev = round(math.sqrt(var))
    all_matched=set()
    for c in candidates: all_matched.update((c.get("analysis") or {}).get("explanation",{}).get("matchedSkills",[]))
    tj = max([((c.get("analysis") or {}).get("explanation",{}).get("totalJDSkills",0)) for c in candidates]+[1])
    cov=round(len(all_matched)/tj*100)
    bias=[]
    if hi>0 and hi/total>0.7: bias.append("High shortlist rate (>70%) — threshold may be too lenient")
    if hi==0: bias.append("No candidates shortlisted — threshold may be too strict")
    if stddev<5 and total>3: bias.append("Low score variance — scoring may not discriminate enough")
    return {"total":total,"avg":avg,"avgATS":avg_ats,"shortlisted":hi,"precision":prec,"recall":rec,"f1":f1,"accuracy":acc,"coverage":cov,"stdDev":stddev,"range":max(scores)-min(scores),"distribution":{"excellent":len([c for c in candidates if c.get("score",0)>=80]),"good":len([c for c in candidates if 60<=c.get("score",0)<80]),"fair":len([c for c in candidates if 40<=c.get("score",0)<60]),"weak":len([c for c in candidates if c.get("score",0)<40])},"confusionMatrix":{"TP":TP,"FP":FP,"FN":FN,"TN":TN},"biasFlags":bias}

# ROUTES
@app.route("/")
def index():
    if not session.get("user"):
        from flask import redirect, url_for
        return redirect(url_for("login_page"))
    return render_template("index.html")

@app.route("/login")
def login_page():
    if session.get("user"):
        from flask import redirect, url_for
        return redirect(url_for("index"))
    return render_template("login.html")

@app.route("/api/health")
def health(): return jsonify({"status":"ok","app":"TalentIQ","pdf_engine":PDF_ENGINE,"docx_engine":DOCX_ENGINE,"features":["resume_parsing","tfidf_matching","ats_scoring","evaluation_metrics","hr_auth"]})

@app.route("/api/login", methods=["POST"])
def api_login():
    d=request.get_json() or {}
    email=(d.get("email") or "").strip().lower(); pw=d.get("password") or ""
    ip = request.headers.get("X-Forwarded-For", request.remote_addr)
    user = resolve_user(email, pw, ip=ip)
    if not user:
        return jsonify({"error":"Invalid email or password"}),401
    session["user"] = user
    return jsonify({"success":True,"user":session["user"]})

@app.route("/api/logout",methods=["POST"])
def api_logout(): session.pop("user",None); return jsonify({"success":True})

@app.route("/api/me")
def api_me():
    u=session.get("user")
    return jsonify({"loggedIn":bool(u),"user":u or {}})

@app.route("/api/parse-resume",methods=["POST"])
def api_parse_resume():
    if "file" not in request.files: return jsonify({"error":"No file"}),400
    f=request.files["file"]
    ext="."+f.filename.rsplit(".",1)[-1].lower() if "." in f.filename else ""
    if ext not in {".pdf",".doc",".docx",".txt"}: return jsonify({"error":f"Unsupported: {ext}"}),400
    try: return jsonify(parse_resume_file(f.read(), f.filename))
    except Exception as e: return jsonify({"error":str(e)}),500

@app.route("/api/parse-resumes-batch",methods=["POST"])
def api_parse_resumes_batch():
    files=request.files.getlist("files")
    if not files: return jsonify({"error":"No files"}),400
    results=[]
    for f in files:
        try: r=parse_resume_file(f.read(),f.filename); results.append({"filename":f.filename,"success":True,**r})
        except Exception as e: results.append({"filename":f.filename,"success":False,"error":str(e)})
    return jsonify({"results":results,"count":len(results)})

@app.route("/api/score",methods=["POST"])
def api_score():
    d=request.get_json() or {}
    rt=d.get("resumeText",""); rn=d.get("resumeName","Candidate"); jdt=d.get("jobDescription","")
    if not rt or not jdt: return jsonify({"error":"resumeText and jobDescription required"}),400
    try: return jsonify(deep_match(rt,rn,jdt))
    except Exception as e: return jsonify({"error":str(e)}),500

@app.route("/api/score-batch",methods=["POST"])
def api_score_batch():
    d=request.get_json() or {}
    resumes=d.get("resumes",[]); jdt=d.get("jobDescription","")
    if not resumes or not jdt: return jsonify({"error":"resumes and jobDescription required"}),400
    results=[]
    for r in resumes:
        try:
            s=deep_match(r.get("text",""),r.get("name","Candidate"),jdt)
            results.append({"name":r.get("name",""),"score":s["match"]["score"],"match":s["match"],"ats":s["ats"]})
        except Exception as e: results.append({"name":r.get("name",""),"error":str(e),"score":0})
    results.sort(key=lambda x:x.get("score",0),reverse=True)
    return jsonify({"ranked":results,"count":len(results)})

@app.route("/api/evaluate",methods=["POST"])
def api_evaluate():
    d=request.get_json() or {}
    cands=d.get("candidates",[])
    if not cands: return jsonify({"error":"candidates required"}),400
    try: return jsonify(compute_evaluation_metrics(cands))
    except Exception as e: return jsonify({"error":str(e)}),500

@app.route("/api/analyze-jd",methods=["POST"])
def api_analyze_jd():
    d=request.get_json() or {}
    jdt=d.get("jobDescription","")
    if not jdt: return jsonify({"error":"jobDescription required"}),400
    return jsonify(analyze_jd(jdt))

@app.route("/api/admin/login-logs", methods=["GET"])
@admin_required
def api_login_logs():
    limit = request.args.get("limit", 100, type=int)
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM login_logs ORDER BY logged_at DESC LIMIT ?", (limit,)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/admin/login-logs/clear", methods=["DELETE"])
@admin_required
def api_clear_logs():
    conn = get_db()
    conn.execute("DELETE FROM login_logs")
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "Login history cleared"})

@app.route("/admin")
def admin_page():
    u = session.get("user")
    if not u or u.get("role") != "admin":
        return redirect(url_for("login_page"))
    return render_template("admin.html")

#  USER MANAGEMENT API (Admin only) 
@app.route("/api/admin/users", methods=["GET"])
@admin_required
def api_get_users():
    conn = get_db()
    rows = conn.execute("SELECT id,name,email,role,created_at,last_login FROM users ORDER BY id").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/admin/users", methods=["POST"])
@admin_required
def api_create_user():
    d = request.get_json() or {}
    name  = (d.get("name") or "").strip()
    email = (d.get("email") or "").strip().lower()
    pw    = (d.get("password") or "").strip()
    role  = (d.get("role") or "hr").strip()
    if not name or not email or not pw:
        return jsonify({"error": "Name, email and password are required"}), 400
    if role not in ("admin", "hr", "recruiter"):
        return jsonify({"error": "Invalid role"}), 400
    try:
        conn = get_db()
        conn.execute("INSERT INTO users (name,email,password,role) VALUES (?,?,?,?)",
                     (name, email, hash_pw(pw), role))
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": f"User {name} created"})
    except sqlite3.IntegrityError:
        return jsonify({"error": "Email already exists"}), 409

@app.route("/api/admin/users/<int:uid>", methods=["PUT"])
@admin_required
def api_update_user(uid):
    d = request.get_json() or {}
    name  = (d.get("name") or "").strip()
    email = (d.get("email") or "").strip().lower()
    role  = (d.get("role") or "hr").strip()
    pw    = (d.get("password") or "").strip()
    if not name or not email:
        return jsonify({"error": "Name and email are required"}), 400
    if role not in ("admin", "hr", "recruiter"):
        return jsonify({"error": "Invalid role"}), 400
    conn = get_db()
    if pw:
        conn.execute("UPDATE users SET name=?,email=?,role=?,password=? WHERE id=?",
                     (name, email, role, hash_pw(pw), uid))
    else:
        conn.execute("UPDATE users SET name=?,email=?,role=? WHERE id=?",
                     (name, email, role, uid))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "User updated"})

@app.route("/api/admin/users/<int:uid>", methods=["DELETE"])
@admin_required
def api_delete_user(uid):
    # Prevent deleting yourself
    me = session.get("user", {})
    if me.get("id") == uid:
        return jsonify({"error": "Cannot delete your own account"}), 400
    conn = get_db()
    conn.execute("DELETE FROM users WHERE id=?", (uid,))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "User deleted"})

if __name__=="__main__":
    init_db()
    print("="*60)
    print("  TalentIQ — Enhanced Resume Screening Suite")
    print("  http://127.0.0.1:5000")
    print(f"  PDF: {PDF_ENGINE or 'none'}  DOCX: {DOCX_ENGINE or 'none'}")
    print("  Default login: admin@talentiq.com / Admin@123")
    print("="*60)
    app.run(debug=True,host="127.0.0.1",port=5000)
